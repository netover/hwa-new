# resync/core/file_ingestor.py
from __future__ import annotations

import logging
import os
import re
import shutil
from pathlib import Path
from typing import Iterator

import docx
import openpyxl
import pypdf  # Corrected import for pypdf
from openpyxl.utils.exceptions import InvalidFileException

from resync.core.exceptions import FileProcessingError, KnowledgeGraphError
from resync.core.interfaces import IFileIngestor, IKnowledgeGraph
from resync.settings import settings

logger = logging.getLogger(__name__)


def is_path_protected(file_path: Path) -> bool:
    """
    Check if a file path is within a protected directory.

    Args:
        file_path: Path to check for protection

    Returns:
        True if the path is within a protected directory, False otherwise
    """
    if not file_path.exists():
        return False

    file_path = file_path.resolve()

    for protected_dir in settings.PROTECTED_DIRECTORIES:
        protected_dir = Path(protected_dir).resolve()
        try:
            # Check if the file path is within the protected directory
            file_path.relative_to(protected_dir)
            return True
        except ValueError:
            # Path is not within the protected directory
            continue

    return False


def is_path_in_knowledge_base(file_path: Path) -> bool:
    """
    Check if a file path is within a knowledge base directory.

    Args:
        file_path: Path to check for knowledge base inclusion

    Returns:
        True if the path is within a knowledge base directory, False otherwise
    """
    if not file_path.exists():
        return False

    file_path = file_path.resolve()

    for knowledge_dir in settings.KNOWLEDGE_BASE_DIRS:
        knowledge_dir = Path(knowledge_dir).resolve()
        try:
            # Check if the file path is within the knowledge base directory
            file_path.relative_to(knowledge_dir)
            return True
        except ValueError:
            # Path is not within the knowledge base directory
            continue

    return False


# --- Text Chunking --- #


def chunk_text(
    text: str, chunk_size: int = 1000, chunk_overlap: int = 200
) -> Iterator[str]:
    """Splits a long text into smaller chunks with overlap."""
    if not text:
        return

    start = 0
    while start < len(text):
        end = start + chunk_size
        yield text[start:end]
        start += chunk_size - chunk_overlap


# --- File Readers --- #


def read_pdf(file_path: Path) -> str:
    """Extracts text from a PDF file."""
    logger.info(f"Reading PDF file: {file_path}")
    try:
        reader = pypdf.PdfReader(file_path)
        text = "".join(
            page.extract_text() for page in reader.pages if page.extract_text()
        )
        return text
    except FileNotFoundError as e:
        logger.error("PDF file not found: %s - %s", file_path, e, exc_info=True)
        return ""
    except PermissionError as e:
        logger.error(
            "Permission denied accessing PDF file: %s - %s", file_path, e, exc_info=True
        )
        return ""
    except pypdf.errors.PdfReadError as e:
        logger.error("PDF read error: %s - %s", file_path, e, exc_info=True)
        return ""
    except ValueError as e:
        logger.error("Invalid PDF content: %s - %s", file_path, e, exc_info=True)
        return ""
    except Exception as e:  # Catch any other pypdf or system errors
        logger.critical(
            "Unexpected error reading PDF %s: %s", file_path, e, exc_info=True
        )
        raise FileProcessingError(f"Failed to process PDF {file_path}") from e


def read_json(file_path: Path) -> str:
    """Extracts text from a JSON file."""
    logger.info(f"Reading JSON file: {file_path}")
    try:
        import json

        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        # Convert JSON to readable text format
        if isinstance(data, dict):
            # For API specifications and structured data, create a readable representation
            text_parts = []
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    text_parts.append(f"{key}: {json.dumps(value, indent=2)}")
                else:
                    text_parts.append(f"{key}: {value}")
            return "\n".join(text_parts)
        elif isinstance(data, list):
            return json.dumps(data, indent=2)
        else:
            return str(data)

    except FileNotFoundError as e:
        logger.error("JSON file not found: %s - %s", file_path, e, exc_info=True)
        return ""
    except PermissionError as e:
        logger.error(
            "Permission denied accessing JSON file: %s - %s",
            file_path,
            e,
            exc_info=True,
        )
        return ""
    except json.JSONDecodeError as e:
        logger.error("Invalid JSON format: %s - %s", file_path, e, exc_info=True)
        return ""
    except UnicodeDecodeError as e:
        logger.error(
            "Encoding error reading JSON file: %s - %s", file_path, e, exc_info=True
        )
        return ""
    except Exception as e:  # Catch other potential OS or parsing errors
        logger.critical(
            "Unexpected error reading JSON %s: %s", file_path, e, exc_info=True
        )
        raise FileProcessingError(f"Failed to process JSON {file_path}") from e


def read_txt(file_path: Path) -> str:
    """Extracts text from a plain text file."""
    logger.info(f"Reading text file: {file_path}")
    try:
        text = file_path.read_text(encoding="utf-8")
        return text
    except FileNotFoundError as e:
        logger.error("Text file not found: %s - %s", file_path, e, exc_info=True)
        return ""
    except PermissionError as e:
        logger.error(
            "Permission denied accessing text file: %s - %s",
            file_path,
            e,
            exc_info=True,
        )
        return ""
    except UnicodeDecodeError as e:
        logger.error(
            "Encoding error reading text file: %s - %s", file_path, e, exc_info=True
        )
        return ""
    except (OSError, IOError) as e:
        logger.critical(
            "Unexpected OS error reading text %s: %s", file_path, e, exc_info=True
        )
        raise FileProcessingError(f"Failed to process text file {file_path}") from e


def read_doc(file_path: Path) -> str:
    """Extracts text from a DOC file (older Word format)."""
    logger.info(f"Reading DOC file: {file_path}")
    try:
        # Try to use python-docx first (it can sometimes handle .doc files)
        try:
            doc = docx.Document(file_path)
            text = "\n".join(para.text for para in doc.paragraphs if para.text)
            if text.strip():
                return text
        except Exception:  # docx library might fail on .doc
            pass

        # Fallback: try to use antiword or similar tool if available
        # For now, return a message indicating the file type is not fully supported
        logger.warning(
            f"DOC file {file_path} requires manual processing or external tool"
        )
        return f"[DOC file: {file_path.name} - Manual processing may be required]"

    except FileNotFoundError as e:
        logger.error("DOC file not found: %s - %s", file_path, e, exc_info=True)
        return ""
    except PermissionError as e:
        logger.error(
            "Permission denied accessing DOC file: %s - %s", file_path, e, exc_info=True
        )
        return ""
    except Exception as e:  # Catch other potential library or system errors
        logger.critical(
            "Unexpected error reading DOC %s: %s", file_path, e, exc_info=True
        )
        raise FileProcessingError(f"Failed to process DOC file {file_path}") from e


def read_xls(file_path: Path) -> str:
    """Extracts text from an XLS file (older Excel format)."""
    logger.info(f"Reading XLS file: {file_path}")
    try:
        # Try to use openpyxl first (it can sometimes handle .xls files)
        try:
            workbook = openpyxl.load_workbook(file_path, read_only=True)
            text_parts = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")

                for row in sheet.iter_rows(values_only=True):
                    # Convert row to text, filtering out None values
                    row_text = " | ".join(str(cell) for cell in row if cell is not None)
                    if row_text.strip():
                        text_parts.append(row_text)

                text_parts.append("")  # Empty line between sheets

            workbook.close()
            return "\n".join(text_parts)

        except Exception:  # openpyxl might fail on .xls
            pass

        # Fallback: try to use xlrd if available, or return message
        try:
            import xlrd

            workbook = xlrd.open_workbook(file_path)
            text_parts = []

            for sheet_name in workbook.sheet_names():
                sheet = workbook.sheet_by_name(sheet_name)
                text_parts.append(f"Sheet: {sheet_name}")

                for row_idx in range(sheet.nrows):
                    row = sheet.row_values(row_idx)
                    # Convert row to text, filtering out None values
                    row_text = " | ".join(
                        str(cell)
                        for cell in row
                        if cell is not None and str(cell).strip()
                    )
                    if row_text.strip():
                        text_parts.append(row_text)

                text_parts.append("")  # Empty line between sheets

            return "\n".join(text_parts)

        except ImportError:
            logger.warning(
                f"XLS file {file_path} requires xlrd library for full support"
            )
            return f"[XLS file: {file_path.name} - Install xlrd for better support]"
        except Exception as e:
            logger.error(f"Error processing XLS file {file_path}: {e}")
            return f"[XLS file: {file_path.name} - Processing error: {e}]"

    except FileNotFoundError as e:
        logger.error("XLS file not found: %s - %s", file_path, e, exc_info=True)
        return ""
    except PermissionError as e:
        logger.error(
            "Permission denied accessing XLS file: %s - %s", file_path, e, exc_info=True
        )
        return ""
    except Exception as e:  # Catch other potential library or system errors
        logger.critical(
            "Unexpected error reading XLS %s: %s", file_path, e, exc_info=True
        )
        raise FileProcessingError(f"Failed to process XLS file {file_path}") from e


def read_md(file_path: Path) -> str:
    """Extracts text from a Markdown file."""
    logger.info(f"Reading Markdown file: {file_path}")
    try:
        text = file_path.read_text(encoding="utf-8")
        return text
    except FileNotFoundError as e:
        logger.error("Markdown file not found: %s - %s", file_path, e, exc_info=True)
        return ""
    except PermissionError as e:
        logger.error(
            "Permission denied accessing Markdown file: %s - %s",
            file_path,
            e,
            exc_info=True,
        )
        return ""
    except UnicodeDecodeError as e:
        logger.error(
            "Encoding error reading Markdown file: %s - %s", file_path, e, exc_info=True
        )
        return ""
    except (OSError, IOError) as e:
        logger.critical(
            "Unexpected OS error reading Markdown %s: %s", file_path, e, exc_info=True
        )
        raise FileProcessingError(f"Failed to process Markdown file {file_path}") from e


def read_docx(file_path: Path) -> str:
    """Extracts text from a DOCX file."""
    logger.info(f"Reading DOCX file: {file_path}")
    try:
        doc = docx.Document(file_path)
        text = "\n".join(para.text for para in doc.paragraphs if para.text)
        return text
    except FileNotFoundError as e:
        logger.error("DOCX file not found: %s - %s", file_path, e, exc_info=True)
        return ""
    except PermissionError as e:
        logger.error(
            "Permission denied accessing DOCX file: %s - %s",
            file_path,
            e,
            exc_info=True,
        )
        return ""  # Corrected import for PackageNotFoundError
    except docx.opc.exceptions.PackageNotFoundError as e:
        logger.error("Invalid DOCX package: %s - %s", file_path, e, exc_info=True)
        return ""
    except ValueError as e:
        logger.error("Invalid DOCX content: %s - %s", file_path, e, exc_info=True)
        return ""
    except Exception as e:  # Catch other potential library or system errors
        logger.critical(
            "Unexpected error reading DOCX %s: %s", file_path, e, exc_info=True
        )
        raise FileProcessingError(f"Failed to process DOCX file {file_path}") from e


def read_excel(file_path: Path) -> str:
    """Extracts text from an XLSX file, iterating through all sheets and cells."""
    logger.info(f"Reading Excel file: {file_path}")
    text_parts = []
    try:
        workbook = openpyxl.load_workbook(file_path, read_only=True)
        for sheetname in workbook.sheetnames:
            sheet = workbook[sheetname]
            text_parts.append(f"Sheet: {sheetname}\n")
            for row in sheet.iter_rows():
                row_texts = []
                for cell in row:
                    if cell.value is not None:
                        row_texts.append(str(cell.value))
                if row_texts:
                    text_parts.append(" | ".join(row_texts))
        return "\n".join(text_parts)
    except FileNotFoundError as e:
        logger.error("Excel file not found: %s - %s", file_path, e, exc_info=True)
        return ""
    except PermissionError as e:
        logger.error(
            "Permission denied accessing Excel file: %s - %s",
            file_path,
            e,
            exc_info=True,
        )
        return ""
    except InvalidFileException as e:
        logger.error("Invalid Excel file: %s - %s", file_path, e, exc_info=True)
        return ""
    except ValueError as e:
        logger.error("Invalid Excel content: %s - %s", file_path, e, exc_info=True)
        return ""
    except Exception as e:  # Catch other potential library or system errors
        logger.critical(
            "Unexpected error reading Excel %s: %s", file_path, e, exc_info=True
        )
        raise FileProcessingError(f"Failed to process Excel file {file_path}") from e


# --- Main Ingestion Logic --- #


class FileIngestor(IFileIngestor):
    """
    Service for ingesting files into the knowledge graph.

    This class handles file uploads, saving, and processing for RAG.
    """

    def __init__(self, knowledge_graph: IKnowledgeGraph):
        """
        Initialize the FileIngestor with dependencies.

        Args:
            knowledge_graph: The knowledge graph service to store extracted content
        """
        self.knowledge_graph = knowledge_graph
        self.rag_directory = settings.BASE_DIR / "rag"
        self.file_readers = {
            ".pdf": read_pdf,
            ".docx": read_docx,
            ".xlsx": read_excel,
            ".md": read_md,
            ".json": read_json,
            ".txt": read_txt,
            ".doc": read_doc,
            ".xls": read_xls,
        }
        # Ensure the RAG directory exists
        self.rag_directory.mkdir(exist_ok=True)
        logger.info(
            f"FileIngestor initialized with RAG directory: {self.rag_directory}"
        )

    async def save_uploaded_file(self, file_name: str, file_content) -> Path:
        """
        Saves an uploaded file to the RAG directory with proper sanitization.

        Args:
            file_name: The original filename
            file_content: A file-like object containing the content

        Returns:
            Path to the saved file

        Raises:
            FileProcessingError: If the file cannot be saved
        """
        # Sanitize the filename to prevent security risks like path traversal
        basename = os.path.basename(file_name)
        # Remove any characters that aren't alphanumeric, dots, underscores, or hyphens
        safe_filename = re.sub(r"[^\w\-_.]", "", basename)
        if not safe_filename:
            raise FileProcessingError("Invalid filename.")

        destination = self.rag_directory / safe_filename

        try:
            logger.info(f"Saving uploaded file: {safe_filename}")
            with destination.open("wb") as buffer:
                shutil.copyfileobj(file_content, buffer)
            logger.info(f"Successfully saved file to {destination}")
            return destination
        except (IOError, OSError) as e:
            logger.critical(
                f"Failed to save uploaded file {safe_filename}: {e}", exc_info=True
            )
            raise FileProcessingError(
                f"Could not save file due to OS error: {e}"
            ) from e

    async def ingest_file(self, file_path: Path) -> bool:
        """
        Ingests a single file into the knowledge graph.

        Args:
            file_path: Path to the file to ingest

        Returns:
            True if ingestion was successful, False otherwise
        """
        if not file_path.exists():
            logger.warning(f"File not found for ingestion: {file_path}")
            return False

        # Check if file is in knowledge base directories
        if not is_path_in_knowledge_base(file_path):
            logger.warning(
                f"File {file_path} is not in knowledge base directories. Skipping ingestion."
            )
            return False

        # Check if file is in protected directories (should not be deleted during processing)
        if is_path_protected(file_path):
            logger.info(
                f"File {file_path} is in protected directory - will not be deleted after processing"
            )

        file_ext = file_path.suffix.lower()
        reader = self.file_readers.get(file_ext)

        if not reader:
            logger.warning(f"Unsupported file type: {file_ext}. Skipping.")
            return False

        # Read the file content
        content = reader(file_path)
        if not content:
            logger.warning(f"No content extracted from {file_path}. Skipping.")
            return False

        # Chunk the content and add to knowledge graph
        chunks = list(chunk_text(content))
        chunk_count = 0
        for i, chunk in enumerate(chunks):
            try:
                metadata = {
                    "source_file": str(file_path.name),
                    "chunk_index": i + 1,
                    "total_chunks": len(chunks),
                }
                # Here we add the chunk to the knowledge graph
                await self.knowledge_graph.add_content(content=chunk, metadata=metadata)
                chunk_count += 1
            except KnowledgeGraphError as e:
                logger.error(
                    "Knowledge graph error adding chunk %d from %s: %s",
                    i + 1,
                    file_path,
                    e,
                    exc_info=True,
                )
            except ValueError as e:
                logger.error(
                    "Value error adding chunk %d from %s: %s",
                    i + 1,
                    file_path,
                    e,
                    exc_info=True,
                )
            except TypeError as e:
                logger.error(
                    "Type error adding chunk %d from %s: %s",
                    i + 1,
                    file_path,
                    e,
                    exc_info=True,
                )
            except Exception:
                logger.critical(
                    "Critical unhandled error adding chunk %d from %s",
                    i + 1,
                    file_path,
                    exc_info=True,
                )
                # We don't re-raise here to allow processing of other chunks

        logger.info(
            f"Successfully ingested {chunk_count}/{len(chunks)} chunks from {file_path}"
        )
        return chunk_count > 0


async def load_existing_rag_documents(file_ingestor: IFileIngestor) -> int:
    """
    Load all existing documents from RAG directories into the knowledge graph.

    Args:
        file_ingestor: The file ingestor instance

    Returns:
        Number of documents processed
    """
    import logging

    logger = logging.getLogger(__name__)
    processed_count = 0

    # Process all knowledge base directories
    for knowledge_dir in settings.KNOWLEDGE_BASE_DIRS:
        knowledge_path = settings.BASE_DIR / knowledge_dir

        if not knowledge_path.exists():
            logger.warning(f"Knowledge base directory not found: {knowledge_path}")
            continue

        logger.info(f"Processing knowledge base directory: {knowledge_path}")

        # Walk through all files in the directory tree
        for file_path in knowledge_path.rglob("*"):
            if file_path.is_file() and not file_path.name.startswith("."):
                # Check if file is in protected directories (should be processed)
                if is_path_in_knowledge_base(file_path):
                    try:
                        logger.info(f"Loading existing document: {file_path.name}")
                        await file_ingestor.ingest_file(file_path)
                        processed_count += 1
                    except FileProcessingError as e:
                        logger.error(
                            "Failed to process document %s: %s",
                            file_path,
                            e,
                            exc_info=True,
                        )
                else:
                    logger.debug(f"Skipping protected file: {file_path}")

    logger.info(f"Loaded {processed_count} existing RAG documents")
    return processed_count


def create_file_ingestor(knowledge_graph: IKnowledgeGraph) -> FileIngestor:
    """
    Factory function to create a FileIngestor instance.

    Args:
        knowledge_graph: The knowledge graph service to use

    Returns:
        A configured FileIngestor instance
    """
    return FileIngestor(knowledge_graph)
